import os
import docx
import json
import nltk
import random
import numpy as np
import pandas as pd
import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers
from keras import models
from keras.layers import Dense, Dropout
import livelossplot
plot_losses = livelossplot.PlotLossesKeras()
import re

from nltk.stem.lancaster import LancasterStemmer
stemmer = LancasterStemmer()
words = []
classes = []
documents = []
ignore_words = [',','.','!','(',')']
json_raw = open('raw.json')
intents = json.load(json_raw)
json_raw.close()

for intent in intents['intents']:
    for pattern in intent['patterns']:
        # tokenize each word in the sentence
        w = nltk.word_tokenize(pattern)
        # add to our words list
        words.extend(w)
        # add to documents in our corpus
        documents.append((w, intent['tag']))
        # add to our classes list
        if intent['tag'] not in classes:
            classes.append(intent['tag'])

words = [stemmer.stem(w.lower()) for w in words if w not in ignore_words]
words = sorted(list(set(words)))
# sort classes
classes = sorted(list(set(classes)))
# documents = combination between patterns and intents
print (len(documents), "documents")
# classes = intents
print (len(classes), "classes", classes)
# words = all words, vocabulary
print (len(words), "unique stemmed words", words)

training = []
# create an empty array for our output
output_empty = [0] * len(classes)
for doc in documents:
    # initialize our bag of words
    bag = []
    # list of tokenized words for the pattern
    pattern_words = doc[0]
    # stem each word - create base word, in attempt to represent related words
    pattern_words = [stemmer.stem(word.lower()) for word in pattern_words]
    # create our bag of words array with 1, if word match found in current pattern
    for w in words:
        bag.append(1) if w in pattern_words else bag.append(0)

    # output is a '0' for each tag and '1' for current tag (for each pattern)
    output_row = list(output_empty)
    output_row[classes.index(doc[1])] = 1

    training.append([bag, output_row])
# shuffle our features and turn into np.array
random.shuffle(training)
training = np.array(training)

train_x = list(training[:, 0])
train_y = list(training[:, 1])
print(train_x)
print(train_y)
x = np.array(train_x)
y = np.array(train_y)
len_x = len(train_x[0])
len_y = len(train_y[0])

model = keras.Sequential([
    keras.layers.Flatten(input_shape=(len_x,)),
    keras.layers.Dense(128, activation=tf.nn.relu),
	keras.layers.Dense(128, activation=tf.nn.relu),
	keras.layers.Dense(128, activation=tf.nn.relu),
    keras.layers.Dense(128, activation=tf.nn.relu),

    keras.layers.Dense(len_y, activation=tf.nn.sigmoid),
])

model.compile(optimizer='adam',
              loss='binary_crossentropy',
              metrics=['accuracy'])

model.fit(train_x, train_y, epochs=50, batch_size=1)

def clean_up_sentence(sentence):
    # tokenize the pattern
    sentence_words = nltk.word_tokenize(sentence)
    # stem each word
    sentence_words = [stemmer.stem(word.lower()) for word in sentence_words  if w not in ignore_words]
    return sentence_words

def bow(sentence, words, show_details=False):
    # tokenize the pattern
    sentence_words = clean_up_sentence(sentence)
    # bag of words
    bag = [0]*len(words)
    for s in sentence_words:
        for i,w in enumerate(words):
            if w == s:
                bag[i] = 1
                if show_details:
                    print ("found in bag: %s" % w)

    return(np.array(bag))
context = {}

ERROR_THRESHOLD = 0.35
def classify(sentence):
    # generate probabilities from the model
    #test
    input_data = pd.DataFrame([bow(sentence, words)], dtype=float, index=['input'])
    results = model.predict([input_data])[0]
    # test end
    #results = model.predict([bow(sentence, words)])[0]
    # filter out predictions below a threshold
    results = [[i, r] for i, r in enumerate(results) if r > ERROR_THRESHOLD]
    #results = [[i,r] for i,r in enumerate(results) if r>ERROR_THRESHOLD]
    # sort by strength of probability
    results.sort(key=lambda x: x[1], reverse=True)
    return_list = []
    for r in results:
        return_list.append((classes[r[0]], r[1]))
    # return tuple of intent and probability
    #print(return_list)
    return return_list

def response(sentence, userID='123', show_details=False):
    results = classify(sentence)
    # if we have a classification then find the matching intent tag
    if results:
        # loop as long as there are matches to process
        while results:
            for i in intents['intents']:
                # find a tag matching the first result
                if i['tag'] == results[0][0]:
                    # set context for this intent if necessary
                    if 'context_set' in i:
                        if show_details: print ('context:', i['context_set'])
                        context[userID] = i['context_set']

                    # check if this intent is contextual and applies to this user's conversation
                    if not 'context_filter' in i or \
                        (userID in context and 'context_filter' in i and i['context_filter'] == context[userID]):
                        if show_details: print ('tag:', i['tag'])
                        # a random response from the intent
                        print(random.choice(i['responses']))
                        return str(random.choice(i['responses']))

            results.pop(0)


list_file = os.listdir('resume')
l = []
responsbility = []
technical = []
experience = []
education = []
positions = []
company_name = []
email_name = []
data = {}
email=[]
name = []
rgx = r'(?:\.?)([\w\-_+#~!$&\'\.]+(?<!\.)(@|[ ]?\(?[ ]?(at|AT)[ ]?\)?[ ]?)(?<!\.)[\w]+[\w\-\.]*\.[a-zA-Z-]{2,3})(?:[^\w])'
for file in list_file:
    data.clear()
    technical.clear()
    responsbility.clear()
    education.clear()
    experience.clear()
    positions.clear()
    email_name.clear()
    company_name.clear()
    email.clear()
    name.clear()
    doc1 = docx.Document("resume/"+file)
    print(" opened resume/"+file)
    l.clear()
    for paras in doc1.paragraphs:
        l.append(paras.text)


    for table in doc1.tables:
        for row in table.rows:
            for cell in row.cells:
                # print(cell.text)
                l.append(cell.text)
    # print(dict)
    print(l)

    for sentences in l:
        #if re.search(rgx,sentences):
         #   match.append(sentences)
        tag = response(sentences)
        if tag == "technical skills":
            technical.append(sentences)
        elif tag == "responsibilities":
            responsbility.append(sentences)
        elif tag == "education":
            education.append(sentences)
        elif tag == "experience":
            experience.append(sentences)
        elif tag == "positions":
            positions.append(sentences)
        elif tag == "company":
            company_name.append(sentences)
        elif tag == 'email_name':
            email_name.append(sentences)
        else:
            pass
    regex1 = r"\"?([-a-zA-Z0-9.`?{}]+@\w+\.\w+)\"?"
    for i in email_name:
        match = re.findall(regex1, str(i))
        if match == []:
            name.append(i)
        else:
            email.append(str(match[0]))


    #for i in email:
    #    index = email_name.index(i)
    #    email_name.remove(index)


    with open("json_files/"+file+".json", 'w+') as json_file:
        data = {'responsibilities':responsbility, 'technological':technical, 'education': education,'experience': experience,'positions': positions,
                'Company_name':company_name, 'email_name': email_name ,'email':email, 'name': name }
        json.dump(data,json_file)



print("end of program")
